from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
OUTPUT = ROOT / "Manual-do-Usuario-Mavi-SFTP.docx"
LOGO = ROOT / "src" / "renderer" / "assets" / "logo.png"

BLUE = "00B8D4"
DARK = "16202A"
MUTED = "5D6873"
LIGHT = "EAF8FB"
PALE = "F4F6F8"
WHITE = "FFFFFF"
RED = "A83232"


def rgb(value):
    return RGBColor.from_string(value)


def font(run, size=11, bold=False, color=DARK, italic=False):
    run.font.name = "Calibri"
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), "Calibri")
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), "Calibri")
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic
    run.font.color.rgb = rgb(color)


def shade(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=90, start=120, bottom=90, end=120):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for name, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{name}"))
        if node is None:
            node = OxmlElement(f"w:{name}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_repeat_table_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    header = OxmlElement("w:tblHeader")
    header.set(qn("w:val"), "true")
    tr_pr.append(header)


def set_table_widths(table, widths):
    table.autofit = False
    for row in table.rows:
        for cell, width in zip(row.cells, widths):
            cell.width = Inches(width)
            set_cell_margins(cell)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def page_field(paragraph):
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = paragraph.add_run("Página ")
    font(run, 9, color=MUTED)
    fld = OxmlElement("w:fldSimple")
    fld.set(qn("w:instr"), "PAGE")
    paragraph._p.append(fld)


def setup_styles(doc):
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(0.8)
    section.bottom_margin = Inches(0.75)
    section.left_margin = Inches(0.85)
    section.right_margin = Inches(0.85)
    section.header_distance = Inches(0.35)
    section.footer_distance = Inches(0.35)

    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    normal.font.size = Pt(10.5)
    normal.font.color.rgb = rgb(DARK)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.16

    for name, size, before, after, color in (
        ("Heading 1", 17, 17, 8, BLUE),
        ("Heading 2", 13.5, 13, 6, BLUE),
        ("Heading 3", 11.5, 9, 4, DARK),
    ):
        style = doc.styles[name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = rgb(color)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True

    for list_style in ("List Bullet", "List Number"):
        style = doc.styles[list_style]
        style.font.name = "Calibri"
        style.font.size = Pt(10.5)
        style.paragraph_format.left_indent = Inches(0.38)
        style.paragraph_format.first_line_indent = Inches(-0.19)
        style.paragraph_format.space_after = Pt(4)
        style.paragraph_format.line_spacing = 1.16


def add_header_footer(section):
    header = section.header
    p = header.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_after = Pt(2)
    r = p.add_run("MAVI SFTP  |  MANUAL DO USUÁRIO")
    font(r, 8.5, bold=True, color=MUTED)
    p_pr = p._p.get_or_add_pPr()
    borders = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "6")
    bottom.set(qn("w:color"), BLUE)
    borders.append(bottom)
    p_pr.append(borders)
    page_field(section.footer.paragraphs[0])


def add_bullet(doc, text, bold_prefix=None):
    p = doc.add_paragraph(style="List Bullet")
    if bold_prefix and text.startswith(bold_prefix):
        a = p.add_run(bold_prefix)
        font(a, 10.5, bold=True)
        b = p.add_run(text[len(bold_prefix):])
        font(b, 10.5)
    else:
        font(p.add_run(text), 10.5)
    return p


def add_step(doc, text):
    p = doc.add_paragraph(style="List Number")
    font(p.add_run(text), 10.5)
    return p


def add_note(doc, label, text, warning=False):
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    cell = table.cell(0, 0)
    cell.width = Inches(6.65)
    shade(cell, "FDEEEE" if warning else LIGHT)
    set_cell_margins(cell, 120, 160, 120, 160)
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(0)
    r = p.add_run(f"{label}: ")
    font(r, 10.3, bold=True, color=RED if warning else DARK)
    font(p.add_run(text), 10.3, color=DARK)
    doc.add_paragraph().paragraph_format.space_after = Pt(1)


def add_feature_table(doc, rows):
    table = doc.add_table(rows=1, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    headers = table.rows[0].cells
    headers[0].text = "Recurso"
    headers[1].text = "O que faz"
    for cell in headers:
        shade(cell, BLUE)
        for run in cell.paragraphs[0].runs:
            font(run, 10, bold=True, color=WHITE)
    set_repeat_table_header(table.rows[0])
    for label, detail in rows:
        cells = table.add_row().cells
        cells[0].text = label
        cells[1].text = detail
        shade(cells[0], PALE)
        for run in cells[0].paragraphs[0].runs:
            font(run, 10, bold=True)
        for run in cells[1].paragraphs[0].runs:
            font(run, 10)
    set_table_widths(table, [1.75, 4.9])
    doc.add_paragraph()


def add_heading(doc, text, level=1):
    return doc.add_heading(text, level=level)


def create_document():
    doc = Document()
    setup_styles(doc)
    section = doc.sections[0]

    # Capa editorial
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(72)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    if LOGO.exists():
        p.add_run().add_picture(str(LOGO), width=Inches(3.7))
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(35)
    p.paragraph_format.space_after = Pt(8)
    font(p.add_run("MANUAL DO USUÁRIO"), 28, bold=True, color=DARK)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(18)
    font(p.add_run("Cliente SFTP para Windows"), 15, color=BLUE)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(70)
    font(p.add_run("Conexão, navegação, transferência e conversão de arquivos"), 11.5, color=MUTED)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(70)
    font(p.add_run("Versão do manual: 1.0  |  Julho de 2026"), 9.5, color=MUTED)
    doc.add_page_break()

    add_header_footer(doc.sections[0])

    add_heading(doc, "Sumário", 1)
    toc = [
        "1. Sobre o Mavi SFTP", "2. Instalação e abertura", "3. Conhecendo a interface",
        "4. Conectar e desconectar", "5. Favoritos", "6. Navegar, pesquisar e ordenar",
        "7. Download de arquivos", "8. Download e conversão de pastas",
        "9. Processamento de Venda Diária", "10. Upload e criação de pastas",
        "11. Renomear e excluir", "12. Acompanhar operações", "13. Boas práticas de segurança",
        "14. Solução de problemas", "15. Referência rápida",
    ]
    for item in toc:
        add_bullet(doc, item)
    add_note(doc, "Objetivo", "Este manual descreve todas as funcionalidades disponíveis na versão atual do Mavi SFTP para Windows.")

    add_heading(doc, "1. Sobre o Mavi SFTP", 1)
    doc.add_paragraph("O Mavi SFTP é um aplicativo para acessar servidores SFTP com segurança, visualizar pastas remotas e transferir arquivos entre o servidor e o computador. Além das operações comuns de arquivos, ele converte arquivos CSV/TXT em Excel e possui um fluxo específico para bases de vendas diárias.")
    add_feature_table(doc, [
        ("Conexão SFTP", "Acesso por host, porta, usuário e senha."),
        ("Navegação", "Visualização de pastas, arquivos, caminho atual e arquivos recentes."),
        ("Transferências", "Download individual, download em lote e upload."),
        ("Conversão", "Conversão de CSV ou TXT para arquivos XLSX."),
        ("Organização", "Pesquisa, ordenação, favoritos, criação, renomeação e exclusão."),
        ("Venda Diária", "Separação por dia, seleção mensal e geração de arquivo concatenado."),
    ])

    add_heading(doc, "2. Instalação e abertura", 1)
    add_step(doc, "Execute o instalador Mavi-SFTP-Setup fornecido pela equipe responsável.")
    add_step(doc, "Escolha a pasta de instalação, quando solicitado, e conclua o assistente.")
    add_step(doc, "Abra o programa pelo atalho da Área de Trabalho ou pelo menu Iniciar.")
    add_note(doc, "Importante", "O instalador completo inclui os componentes usados pelo aplicativo. Não é necessário instalar Node.js, Electron ou Python separadamente.")
    add_heading(doc, "2.1 Aviso do Microsoft Defender SmartScreen", 2)
    doc.add_paragraph("Como o instalador atual ainda não possui uma assinatura digital reconhecida, o Windows pode exibir a mensagem “O Windows protegeu o computador” e informar que o Microsoft Defender SmartScreen impediu a execução de um aplicativo não reconhecido.")
    add_note(doc, "O que significa", "Esse aviso indica que o Windows ainda não conseguiu confirmar a identidade do publicador ou estabelecer reputação para aquela versão do arquivo. Ele não significa, por si só, que um vírus foi detectado.", warning=True)
    doc.add_paragraph("Prossiga somente quando o instalador tiver sido obtido na Release oficial do repositório do Mavi SFTP ou enviado diretamente por uma pessoa responsável pelo aplicativo.")
    add_step(doc, "Na tela do SmartScreen, confirme que o arquivo aberto é o instalador esperado do Mavi SFTP.")
    add_step(doc, "Clique em Mais informações.")
    add_step(doc, "Verifique novamente o nome do aplicativo e a origem do arquivo.")
    add_step(doc, "Clique em Executar assim mesmo para iniciar a instalação.")
    add_note(doc, "Não prossiga", "Se o arquivo veio de um link desconhecido, mensagem suspeita, site não oficial ou fonte que você não consegue confirmar, cancele a instalação e solicite um novo arquivo à equipe responsável.", warning=True)
    doc.add_paragraph("O aviso pode reaparecer em novas versões porque cada instalador possui um arquivo e um hash diferentes. A solução definitiva é assinar digitalmente o aplicativo com um certificado de assinatura de código confiável. Até que essa assinatura seja configurada, o aviso é esperado em computadores que ainda não reconhecem o instalador.")

    add_heading(doc, "3. Conhecendo a interface", 1)
    add_feature_table(doc, [
        ("Barra superior", "Exibe a marca, o estado da conexão e os botões minimizar, maximizar/restaurar e fechar."),
        ("Conexão", "Campos Host, Porta, Usuário e Senha, além de Conectar/Desconectar."),
        ("Favoritos", "Lista conexões salvas e permite preencher rapidamente host, porta e usuário."),
        ("Log", "Registra conexões, transferências, conversões, avisos e erros."),
        ("Barra de caminho", "Mostra a pasta atual e reúne navegação, pesquisa e ações da pasta."),
        ("Lista de arquivos", "Exibe nome, tamanho, data de modificação e ações disponíveis."),
        ("Progresso", "Mostra percentual, quantidade processada, nome do arquivo e erros."),
    ])

    add_heading(doc, "4. Conectar e desconectar", 1)
    add_heading(doc, "4.1 Fazer uma conexão", 2)
    add_step(doc, "No campo Host, informe o endereço do servidor, sem incluir sftp://.")
    add_step(doc, "Informe a Porta. O valor padrão de SFTP é 22.")
    add_step(doc, "Preencha Usuário e Senha.")
    add_step(doc, "Clique em Conectar ou pressione Enter em um dos campos.")
    add_step(doc, "Aguarde a confirmação. Após conectar, o aplicativo abre a pasta raiz (/).")
    add_note(doc, "Validação", "Host e usuário são obrigatórios. O tempo máximo inicial de resposta do servidor é de aproximadamente 10 segundos.")
    add_heading(doc, "4.2 Identificar o estado", 2)
    add_bullet(doc, "Desconectado: a área principal mostra que não há conexão ativa.")
    add_bullet(doc, "Conectado: a barra superior exibe usuário@host e o navegador de arquivos fica disponível.")
    add_heading(doc, "4.3 Desconectar", 2)
    doc.add_paragraph("Clique em Desconectar antes de trocar de servidor ou encerrar uma sessão. Ao fechar o aplicativo, a conexão ativa também é finalizada.")

    add_heading(doc, "5. Favoritos", 1)
    add_heading(doc, "5.1 Salvar", 2)
    add_step(doc, "Conecte-se ao servidor desejado.")
    add_step(doc, "Clique em + Salvar conexão atual.")
    add_step(doc, "Digite um nome que identifique a conexão e confirme.")
    add_note(doc, "Privacidade", "O favorito guarda apenas nome, host, porta e usuário no computador. A senha não é salva.")
    add_heading(doc, "5.2 Usar ou remover", 2)
    add_bullet(doc, "Clique no nome do favorito para preencher Host, Porta e Usuário; depois informe a senha e conecte.")
    add_bullet(doc, "Clique no X ao lado do favorito para removê-lo.")

    add_heading(doc, "6. Navegar, pesquisar e ordenar", 1)
    add_heading(doc, "6.1 Navegação", 2)
    add_bullet(doc, "Abra uma pasta com duplo clique.")
    add_bullet(doc, "Use a seta para cima para retornar à pasta anterior.")
    add_bullet(doc, "Consulte o caminho completo na barra de caminho.")
    add_bullet(doc, "Clique em Atualizar para recarregar a pasta atual.")
    add_heading(doc, "6.2 Arquivos recentes", 2)
    doc.add_paragraph("Na pasta raiz, a seção Recentes lista até 30 arquivos do servidor, ordenados pela data de modificação mais recente. O caminho completo aparece no nome para indicar onde cada arquivo está armazenado.")
    add_heading(doc, "6.3 Pesquisa e ordenação", 2)
    add_bullet(doc, "Digite no campo Buscar arquivo para filtrar os itens da pasta atual pelo nome.")
    add_bullet(doc, "Clique em Nome, Tamanho ou Modificado para ordenar.")
    add_bullet(doc, "Clique novamente no mesmo cabeçalho para alternar entre crescente e decrescente.")
    add_bullet(doc, "As pastas permanecem antes dos arquivos, independentemente da ordenação escolhida.")

    add_heading(doc, "7. Download de arquivos", 1)
    add_heading(doc, "7.1 Baixar no formato original", 2)
    add_step(doc, "Localize o arquivo desejado.")
    add_step(doc, "Clique em down na linha do arquivo.")
    add_step(doc, "Escolha o nome e o local de destino no computador.")
    add_step(doc, "Acompanhe o progresso e aguarde a mensagem de conclusão.")
    add_heading(doc, "7.2 Baixar CSV/TXT como Excel", 2)
    add_step(doc, "Em um arquivo CSV ou TXT, clique em xlsx.")
    add_step(doc, "Escolha onde salvar o arquivo .xlsx.")
    add_step(doc, "Aguarde a conversão e confira os avisos no Log.")
    add_note(doc, "Conversão", "O aplicativo detecta codificação UTF-8 ou Latin-1 e separadores como ponto e vírgula, vírgula, tabulação e barra vertical. Colunas identificadoras, como código, EAN, CPF e CNPJ, são preservadas como texto.")
    add_note(doc, "Planilhas grandes", "Uma base muito grande pode ser dividida em mais de um arquivo Excel para respeitar limites seguros de linhas.")

    add_heading(doc, "8. Download e conversão de pastas", 1)
    doc.add_paragraph("O download em lote pode ser iniciado pelos botões da barra superior da pasta atual ou pelos botões csv, xlsx e ambos exibidos em uma pasta que contém arquivos compatíveis.")
    add_feature_table(doc, [
        ("CSV", "Salva os arquivos CSV/TXT selecionados mantendo o formato original."),
        ("xlsx", "Converte os arquivos CSV/TXT encontrados para Excel."),
        ("Ambos", "Salva uma cópia original e outra em Excel para cada arquivo."),
    ])
    add_heading(doc, "8.1 Escolher o conteúdo", 2)
    doc.add_paragraph("Após iniciar o download, a janela Escolher download apresenta opções conforme os nomes existentes na pasta:")
    add_bullet(doc, "Baixar Todos Arquivos: processa todos os CSV/TXT da pasta.")
    add_bullet(doc, "Baixar período MM/AAAA: processa somente arquivos cujo nome contém uma data válida no padrão AAAAMMDD para o período escolhido.")
    add_bullet(doc, "Baixar Venda Diária do mês MM/AAAA: inicia o tratamento especial descrito na próxima seção.")
    add_step(doc, "Selecione a opção desejada.")
    add_step(doc, "Escolha uma pasta local de destino.")
    add_step(doc, "Aguarde o processamento terminar antes de fechar o aplicativo.")

    add_heading(doc, "9. Processamento de Venda Diária", 1)
    doc.add_paragraph("Este modo reconhece arquivos com os padrões BR_VENTAS_AAAAMMDD e BR_VENDAS_DIARIA_COCACOLA_GPA_ENERGETICOS_SEM_CONCORRENCIA_AAAAMMDD, com ou sem extensão CSV/TXT.")
    add_heading(doc, "9.1 Como o período é calculado", 2)
    add_bullet(doc, "O aplicativo usa arquivos desde o primeiro dia do mês selecionado até dez dias depois do último dia do mês.")
    add_bullet(doc, "Dentro de cada base, utiliza a coluna FECHA_COMERCIAL para separar os registros por dia.")
    add_bullet(doc, "Somente datas no formato AAAAMMDD pertencentes ao mês escolhido entram nos arquivos finais.")
    add_heading(doc, "9.2 Arquivos gerados", 2)
    add_bullet(doc, "Um arquivo por dia encontrado, no formato solicitado (CSV, XLSX ou ambos).")
    add_bullet(doc, "Um arquivo CONCATENADO_AAAAMM, reunindo os dias finais em ordem.")
    add_bullet(doc, "Quando houver mais de uma fonte para o mesmo dia, o processamento mantém a versão mais recente tratada para aquele dia.")
    add_note(doc, "Requisito", "A base precisa possuir a coluna FECHA_COMERCIAL. Se ela não existir, o arquivo será informado como erro no Log.", warning=True)

    add_heading(doc, "10. Upload e criação de pastas", 1)
    add_heading(doc, "10.1 Enviar um arquivo", 2)
    add_step(doc, "Navegue até a pasta remota de destino.")
    add_step(doc, "Clique em Upload.")
    add_step(doc, "Selecione um arquivo do computador.")
    add_step(doc, "Aguarde a confirmação e confira o item na lista atualizada.")
    add_heading(doc, "10.2 Criar uma pasta", 2)
    add_step(doc, "Navegue até a pasta que receberá a nova subpasta.")
    add_step(doc, "Clique em + Pasta.")
    add_step(doc, "Digite o nome e clique em Confirmar.")

    add_heading(doc, "11. Renomear e excluir", 1)
    add_heading(doc, "11.1 Renomear", 2)
    add_step(doc, "Clique em ren na linha do arquivo ou pasta.")
    add_step(doc, "Informe o novo nome e confirme.")
    add_heading(doc, "11.2 Excluir", 2)
    add_step(doc, "Clique em del na linha do item.")
    add_step(doc, "Digite exatamente o nome atual do arquivo ou pasta.")
    add_step(doc, "Confirme a operação.")
    add_note(doc, "Atenção", "A exclusão é realizada diretamente no servidor. Pastas são removidas recursivamente, inclusive com seu conteúdo. Não há lixeira ou botão de desfazer.", warning=True)

    add_heading(doc, "12. Acompanhar operações", 1)
    add_feature_table(doc, [
        ("Indicador superior", "Mostra se o aplicativo está conectado e identifica usuário e servidor."),
        ("Notificações", "Mensagens rápidas confirmam sucesso, informação ou erro."),
        ("Log", "Mantém o histórico textual da sessão, incluindo caminhos, avisos e erros."),
        ("Barra de progresso", "Apresenta percentual, arquivos concluídos e quantidade de erros."),
    ])
    doc.add_paragraph("Em downloads em lote, um erro em determinado arquivo não impede necessariamente os demais. Ao final, confira a quantidade de erros e os detalhes registrados no Log.")

    add_heading(doc, "13. Boas práticas de segurança", 1)
    add_bullet(doc, "Confirme host, pasta e nome do item antes de enviar, renomear ou excluir.")
    add_bullet(doc, "Não compartilhe senhas nem inclua credenciais em nomes de favoritos.")
    add_bullet(doc, "Use contas com apenas as permissões necessárias para o trabalho.")
    add_bullet(doc, "Desconecte-se quando terminar, especialmente em computadores compartilhados.")
    add_bullet(doc, "Evite fechar o programa durante downloads, uploads ou conversões.")
    add_bullet(doc, "Valide os arquivos gerados antes de substituir bases de produção.")

    add_heading(doc, "14. Solução de problemas", 1)
    add_feature_table(doc, [
        ("Preencha host e usuário", "Informe os dois campos obrigatórios antes de conectar."),
        ("Falha na conexão", "Confira host, porta, usuário, senha, internet/VPN e liberação do servidor."),
        ("Tempo esgotado", "O servidor não respondeu no prazo; teste a rede ou solicite validação à infraestrutura."),
        ("Não conectado", "Reconecte antes de tentar listar ou transferir arquivos."),
        ("Nenhum CSV encontrado", "A pasta ou o período escolhido não possui arquivos .csv ou .txt compatíveis."),
        ("Coluna não encontrada", "No modo Venda Diária, confirme a presença de FECHA_COMERCIAL no cabeçalho."),
        ("Erro ao converter", "Verifique espaço em disco, acesso à pasta de destino e integridade do arquivo de origem."),
        ("Nome não confere", "Na exclusão, digite o nome completo exatamente como aparece na lista."),
        ("Ícone/lista não atualizou", "Clique em Atualizar ou feche e abra novamente o aplicativo."),
    ])

    add_heading(doc, "15. Referência rápida", 1)
    add_feature_table(doc, [
        ("↑", "Voltar para a pasta anterior."),
        ("Upload", "Enviar um arquivo local para a pasta remota atual."),
        ("+ Pasta", "Criar uma subpasta na pasta atual."),
        ("↓ CSV", "Baixar em lote no formato original."),
        ("↓ xlsx", "Baixar em lote convertendo para Excel."),
        ("↓ Ambos", "Baixar original e Excel."),
        ("Atualizar", "Recarregar a listagem da pasta atual."),
        ("down", "Baixar um arquivo individual."),
        ("xlsx", "Baixar um CSV/TXT individual como Excel."),
        ("ren", "Renomear arquivo ou pasta."),
        ("del", "Excluir arquivo ou pasta mediante confirmação do nome."),
    ])
    add_note(doc, "Suporte", "Ao solicitar ajuda, informe a operação realizada, a mensagem exibida e as últimas linhas do painel Log. Nunca envie sua senha.")

    # Mantém títulos com seu primeiro conteúdo e adiciona metadados básicos.
    doc.core_properties.title = "Manual do Usuário - Mavi SFTP"
    doc.core_properties.subject = "Guia completo de funcionalidades do cliente Mavi SFTP"
    doc.core_properties.author = "Mavi"
    doc.core_properties.keywords = "Mavi, SFTP, manual, CSV, XLSX, vendas diárias"
    doc.save(OUTPUT)
    print(OUTPUT)


if __name__ == "__main__":
    create_document()
